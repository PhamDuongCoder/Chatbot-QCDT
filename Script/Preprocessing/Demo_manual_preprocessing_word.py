#this file indicates how someone should preprocess a word file manually
#the example is on this file: C:\Users\LENOVO\Documents\GitHub\Chatbot QCĐT\Data\Hoc_phi\[your_file].docx

from docx import Document
import pandas as pd
import re

# defining preprocessing functions

def basic_clean(text):
    """cleaning the text"""
    #delete page number
    text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
    #delete blank line
    text = re.sub(r'\n{3,}', '\n\n', text)
    #delete blank space at the beginning or end of a line
    lines = [line.strip() for line in text.split('\n')]
    text = ('\n').join(lines)
    #normalize vietnamese syntax 
    text = text.replace('\u2013', '-')   # en dash → gạch ngang
    text = text.replace('\u2019', "'")   # curly quote

    return text.strip()

def dataframe_to_markdown(df):
    """Convert pandas DataFrame to markdown tables"""
    #create header row
    headers = [str(col).strip() for col in df.columns]
    md_table = "| " + " | ".join(headers) + " |\n"

    #Create separator row
    md_table += "|" + "|".join(["---" for _ in headers]) + "|\n"

    #Add data rows
    for _, row in df.iterrows():
        values = [str(v).strip() for v in row.values if str(v).strip() != 'nan']
        if values: # Only add if row has content
            md_table += "| " + " | ".join(values) + " |\n"

    return md_table.strip()

def table_to_dataframe(table):
    """Convert Word table to pandas DataFrame"""
    data = []
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        data.append(row_data)
    
    # Use first row as header if not empty
    if data:
        df = pd.DataFrame(data[1:], columns=data[0])
        return df
    return None

def extract_text_and_tables(path):
    """Extract text and tables from Word document"""
    doc = Document(path)
    content_data = {
        "text": "",
        "tables": []
    }
    
    # Create a map of table elements for quick lookup
    table_elements = set()
    for table in doc.tables:
        table_elements.add(table._element)
    
    # Iterate through body elements in order
    for element in doc.element.body:
        if element.tag.endswith('}p'):
            # This is a paragraph - find corresponding paragraph object
            for para in doc.paragraphs:
                if para._element == element:
                    content_data["text"] += para.text + "\n"
                    break
        elif element.tag.endswith('}tbl'):
            # This is a table - find corresponding table object
            for table in doc.tables:
                if table._element == element:
                    df = table_to_dataframe(table)
                    if df is not None:
                        content_data["tables"].append(df)
                    break
    
    return content_data

def merge_text_and_tables(content_data, source_file, category, year):
    """Merge cleaned text and tables into one document"""
    header = f"""
    Nguồn: {source_file}
    Danh mục: {category}
    Năm: {year}
    Ngày xử lý: 8/4/2026

---NỘI DUNG---

"""
    
    text = basic_clean(content_data["text"])
    body = f"\n{text}\n"
    
    #add cleaned tables
    if content_data["tables"]:
        body += "\n--- BẢNG ---\n"
        for df in content_data["tables"]:
            md_table = dataframe_to_markdown(df)
            body += f"\n{md_table}\n"
    
    return header + body

def save_clean_txt(text, output_path):
    """save cleaned text"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

# Preprocessing data (example usage)
# content_data = extract_text_and_tables('.\Data\Hoc_phi\your_file.docx')
# merged = merge_text_and_tables(
#     content_data,
#     source_file='your_file.docx',
#     category='Hoc_phi',
#     year='2025-2026'
# )
# save_clean_txt(merged, "Data/Hoc_phi/your_file.txt")
